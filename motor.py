import re
import io
import os
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TAXA_COMISSAO = 0.043
BASE_DIR      = os.path.dirname(os.path.abspath(__file__))

# ───────── FORMATAR ─────────────────────────────────────────────────────────

def formatar(valor):
    return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def formatar_valor(valor):          # alias usado pelo app.py
    return formatar(valor)

def formatar_percentual(valor):
    return f"{valor:.2f}".replace(".", ",") + "%"

def extenso(valor):
    return f"{formatar(valor)} reais"

def _iso_to_br(date_str):
    if not date_str or len(date_str) != 10:
        return date_str
    try:
        y, m, d = date_str.split('-')
        return f"{d}/{m}/{y}"
    except Exception:
        return date_str

def formatar_telefone(tel):
    digits = re.sub(r'\D', '', tel or '')
    if digits.startswith('55') and len(digits) > 11:
        digits = digits[2:]
    if len(digits) == 11:
        return f"({digits[:2]}) {digits[2:7]}-{digits[7:]}"
    if len(digits) == 10:
        return f"({digits[:2]}) {digits[2:6]}-{digits[6:]}"
    return tel

# ───────── COMISSÃO ─────────────────────────────────────────────────────────

def definir_tipo_comissao(preco, sinal, p30=0, p60=0):
    if not preco or preco <= 0:
        return {"tipo": "faturada", "total_comissao": 0,
                "valor_venda_ajustado": preco or 0, "parcela_desconto": None}

    comissao = round(preco * TAXA_COMISSAO, 2)
    limiar   = preco * 0.10

    if sinal >= limiar:
        tipo, parcela_desconto = "destacada", "ato"
    elif sinal + p30 >= limiar:
        tipo, parcela_desconto = "destacada", "complemento_30"
    elif sinal + p30 + p60 >= limiar:
        tipo, parcela_desconto = "destacada", "complemento_60"
    else:
        tipo, parcela_desconto = "faturada", None

    valor_ajustado = round(preco - comissao, 2) if tipo == "destacada" else preco

    return {
        "tipo":                tipo,
        "total_comissao":      comissao,
        "valor_venda_ajustado": valor_ajustado,
        "parcela_desconto":    parcela_desconto,
    }

# ───────── INTERPRETAR FLUXO ────────────────────────────────────────────────

def interpretar_fluxo(texto):
    parcelas    = []
    valor_venda = 0

    for linha in (texto or "").split("\n"):
        linha = linha.strip()
        if not linha:
            continue

        linha_low = linha.lower()

        if "valor de venda" in linha_low:
            v = re.search(r"([\d\.,]+)", linha)
            if v:
                valor_venda = float(v.group(1).replace(".", "").replace(",", "."))

        valor_match = re.search(r":\s*([\d\.,]+)", linha)
        valor = float(valor_match.group(1).replace(".", "").replace(",", ".")) if valor_match else 0

        qtd_match = re.search(r"\((\d+)\)", linha)
        qtd = int(qtd_match.group(1)) if qtd_match else 1

        data_match = re.search(r"(\d{2}/\d{2}/\d{4})", linha)
        data = data_match.group(1) if data_match else ""

        if   "ato"          in linha_low:                          tipo = "ato"
        elif "complemento"  in linha_low:                          tipo = "complemento"
        elif "mensais"      in linha_low:                          tipo = "mensal"
        elif "semestrais"   in linha_low:                          tipo = "semestral"
        elif "anuais"       in linha_low:                          tipo = "anual"
        elif "financiamento" in linha_low:                         tipo = "financiamento"
        elif "única" in linha_low or "unica" in linha_low:        tipo = "unica"
        else: continue

        parcelas.append({"tipo": tipo, "qtd": qtd, "valor": valor,
                         "data": data, "reajustavel": True})

    return parcelas, valor_venda

# ───────── PAGAMENTO ────────────────────────────────────────────────────────

def romano(n):
    r = ["", "i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x"]
    return r[n] if n < len(r) else str(n)

def montar_pagamento(parcelas):
    textos = []
    i = 1

    for p in parcelas:
        if p["tipo"] == "ato":
            continue

        total      = p["valor"] * p["qtd"]
        valor_unit = formatar(p["valor"])
        valor_tot  = formatar(total)
        qtd        = p["qtd"]
        data       = p.get("data", "")
        reaj       = "reajustável" if p.get("reajustavel", True) else "fixo"

        if p["tipo"] in ("financiamento", "unica"):
            label = "Parcela de Financiamento" if p["tipo"] == "financiamento" else "Parcela Única"
            textos.append(
                f"({romano(i)}) R${valor_unit} ({extenso(p['valor'])}) serão pagos em uma única "
                f"parcela {reaj} com vencimento em {data} (\"{label}\");"
            )
            i += 1
            continue

        if   p["tipo"] == "mensal":      descricao, period = "Parcelas Mensais", "mensais"
        elif p["tipo"] == "anual":       descricao, period = "Parcelas Anuais", "anuais"
        elif p["tipo"] == "semestral":   descricao, period = "Parcelas Semestrais", "semestrais"
        elif p["tipo"] == "complemento": descricao, period = "Parcelas de Complemento de Sinal", "mensais"
        else:                            descricao, period = "Parcelas", "mensais"

        textos.append(
            f"({romano(i)}) R${valor_tot} ({extenso(total)}) serão pagos em {qtd} parcelas "
            f"{reaj}s, {period}, sucessivas, no valor de R${valor_unit} ({extenso(p['valor'])}) "
            f"cada uma delas, vencendo-se a primeira no dia {data} e as demais no mesmo dia dos "
            f"{period} subsequentes (\"{descricao}\");"
        )
        i += 1

    return "\n\n".join(textos)

# ───────── COMPRADOR ────────────────────────────────────────────────────────

def qualificar(c):
    inscrito = "inscrito" if c.get("sexo", "M") == "M" else "inscrita"
    rg_part  = (f", RG nº {c['rg']} expedido pelo {c.get('orgao_emissor', '')}"
                if c.get("rg") else "")
    return (
        f"{c['nome']}, {c.get('nacionalidade', 'brasileiro(a)').lower()}, "
        f"{c.get('profissao', '')}, {inscrito} no CPF/ME sob o nº {c['cpf']}{rg_part}"
    )

def _qualificar_simples(c):
    return (
        f"{c['nome']}, {c.get('nacionalidade', 'brasileiro(a)').lower()}, "
        f"{c.get('profissao', '')}"
    )

def _rg_str(c):
    if not c.get("rg"):
        return None
    s = c["rg"]
    if c.get("orgao_emissor"):
        s += f" do {c['orgao_emissor']}"
    if c.get("data_emissao"):
        s += f" em {_iso_to_br(c['data_emissao'])}"
    return s

def montar_compradora(dados):
    compradores = dados["compradores"]
    relacao     = dados.get("relacao", "solteiro / independentes")
    regime      = dados.get("regime_bens", "")
    endereco    = dados.get("endereco", "")

    if len(compradores) == 1:
        return f"{qualificar(compradores[0])}, residente(s) na {endereco}"

    c1, c2 = compradores[0], compradores[1]

    if relacao == "casado":
        conj = "sua mulher" if c2.get("sexo") == "F" else "seu marido"

        cpf_part = f"inscritos no CPF/ME sob os nºs {c1['cpf']} e {c2['cpf']}"

        rg1, rg2 = _rg_str(c1), _rg_str(c2)
        if rg1 and rg2:
            rg_part = f", portadores das identidades nºs {rg1} e {rg2}"
        elif rg1:
            rg_part = f", portador(a) da identidade nº {rg1}"
        elif rg2:
            rg_part = f", portador(a) da identidade nº {rg2}"
        else:
            rg_part = ""

        return (
            f"{_qualificar_simples(c1)}, e {conj} {_qualificar_simples(c2)}, "
            f"casados pelo regime da {regime.lower()}, "
            f"{cpf_part}{rg_part}, "
            f"residente(s) na {endereco}"
        )

    if relacao == "união estável":
        data_esc  = dados.get("data_escritura", "")
        data_part = f" desde {data_esc}" if data_esc else ""
        return (
            f"{qualificar(c1)}, e {qualificar(c2)}, "
            f"companheiros em união estável{data_part}, residentes na {endereco}"
        )

    return f"{qualificar(c1)}, e {qualificar(c2)}, residentes na {endereco}"

# ───────── SUBSTITUIÇÃO (preserva formatação dos runs) ──────────────────────

def _substituir_runs(paragrafo, subs):
    for run in paragrafo.runs:
        for placeholder, valor in subs.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(valor))

def _substituir_doc(doc, subs):
    for p in doc.paragraphs:
        _substituir_runs(p, subs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _substituir_runs(p, subs)

# ───────── MESCLAR DOCUMENTOS ───────────────────────────────────────────────

def _mesclar(doc_base, doc_extra):
    """Adiciona o conteúdo de doc_extra ao final de doc_base."""
    for element in doc_extra.element.body:
        doc_base.element.body.append(copy.deepcopy(element))

def _force_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)

# ───────── NÚCLEO DE GERAÇÃO ────────────────────────────────────────────────

def _gerar_buffer(dados: dict) -> io.BytesIO:
    # Parcelas
    if dados.get("parcelas_diretas"):
        parcelas = [dict(p) for p in dados["parcelas_diretas"]]
        preco    = float(dados.get("preco_direto") or 0)
    else:
        parcelas, preco_fluxo = interpretar_fluxo(dados.get("fluxo", ""))
        preco = float(dados.get("valor_venda_manual") or 0) or preco_fluxo

    sinal      = next((p["valor"] for p in parcelas if p["tipo"] == "ato"), 0.0)
    percentual = (sinal / preco * 100) if preco > 0 else 0

    # Tipo de comissão (para escolher o par de templates)
    comps      = [p for p in parcelas if p["tipo"] == "complemento"]
    p30        = comps[0]["valor"] if comps else 0.0
    p60        = comps[1]["valor"] if len(comps) >= 2 else 0.0
    comissao   = definir_tipo_comissao(preco, sinal, p30, p60)

    # Override manual da parcela de desconto
    if dados.get("parcela_desconto_manual"):
        comissao["tipo"] = "destacada"
        comissao["parcela_desconto"] = dados["parcela_desconto_manual"]

    tipo_comissao = comissao["tipo"]

    # Escolher templates
    if tipo_comissao == "destacada":
        f_contrato = os.path.join(BASE_DIR, "contrato_cabeca.docx")
        f_corpo    = os.path.join(BASE_DIR, "corpo_cabeça.docx")
    else:
        f_contrato = os.path.join(BASE_DIR, "contrato_faturado.docx")
        f_corpo    = os.path.join(BASE_DIR, "corpo_faturado.docx")

    doc_contrato = Document(f_contrato)
    doc_corpo    = Document(f_corpo)

    # Imobiliária (faturado) ou vazio (destacada)
    imob_str = ""
    if tipo_comissao == "faturada" and dados.get("imobiliarias"):
        im = dados["imobiliarias"][0]
        imob_str = im.get("empresa", "") or im.get("nome", "")

    # Comunicação = telefone + e-mail do comprador
    tel    = formatar_telefone(dados.get("telefone", ""))
    email  = dados.get("email", "")
    comunicacao = " | ".join(filter(None, [tel, email]))

    # Tipo de assinatura legível
    tipo_ass_raw = dados.get("tipo_ass", "digital")
    tipo_ass_str = "meio digital" if tipo_ass_raw == "digital" else "2 (duas) vias físicas de igual forma e teor"

    # Nomes para bloco de assinatura
    compradores = dados.get("compradores", [{}])
    ass1 = compradores[0].get("nome", "") if compradores else ""
    ass2 = compradores[1].get("nome", "") if len(compradores) > 1 else ""

    subs_contrato = {
        "«COMPRADORA»":       montar_compradora(dados),
        "«COMUNICACAO»":      comunicacao,
        "«FRACAO_IDEAL»":     dados.get("fracao_ideal", ""),
        "«IMOBILIARIA»":      imob_str,
        "«PAGAMENTO»":        montar_pagamento(parcelas),
        "«PORCENTAGEMSINAL»": formatar_percentual(percentual),
        "«PRECO»":            formatar(preco),
        "«SINAL»":            formatar(sinal),
        "«UNIDADE»":          dados.get("unidade", ""),
        "«VAGAS»":            str(dados.get("vagas") or ""),
        "«VLR_COMISSAO»":     formatar(comissao["total_comissao"]),
    }

    subs_corpo = {
        "«ASS_1»":           ass1,
        "«ASS_2»":           ass2,
        "«Data_Assinatura»": dados.get("data_assinatura", ""),
        "«TIPO_ASS»":        tipo_ass_str,
        "«UNIDADE»":         dados.get("unidade", ""),
    }

    _substituir_doc(doc_contrato, subs_contrato)
    _substituir_doc(doc_corpo,    subs_corpo)

    for p in doc_corpo.paragraphs:
        if p.text.strip().startswith('E por assim se acharem'):
            _force_page_break_before(p)
            break

    _mesclar(doc_contrato, doc_corpo)

    buf = io.BytesIO()
    doc_contrato.save(buf)
    buf.seek(0)
    return buf

# ───────── API PÚBLICA ───────────────────────────────────────────────────────

def gerar_contrato_bytes(dados: dict) -> bytes:
    """Retorna o .docx como bytes — usado pela API FastAPI."""
    return _gerar_buffer(dados).read()

def gerar_contrato(dados: dict) -> str:
    """Salva em disco e retorna caminho — compatível com app.py Streamlit."""
    buf  = _gerar_buffer(dados)
    path = "contrato_gerado.docx"
    with open(path, "wb") as f:
        f.write(buf.read())
    return path
